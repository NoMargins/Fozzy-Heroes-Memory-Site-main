from docx import Document
import json


# Load the document
doc_path = "../heroes/allDesks.docx"
doc = Document(doc_path)


# Load the document
doc = Document(doc_path)

# Extract paragraphs and process
profiles = []
current_profile = None

for para in doc.paragraphs:
    text = para.text.strip()
    if text:
        if text.startswith("name:"):
            if current_profile:
                profiles.append(current_profile)
            current_profile = {"name": text[5:].strip(), "description": ""}
        elif current_profile:
            if current_profile["description"]:
                current_profile["description"] += "\n\n" + text
            else:
                current_profile["description"] = text

# Append the last profile if exists
if current_profile:
    profiles.append(current_profile)

# Convert to JSON
json_output = json.dumps(profiles, ensure_ascii=False, indent=2)

# Save JSON to a file
with open("profiles.json", "w", encoding='utf-8') as f:
    f.write(json_output)

print("Profiles have been extracted and saved to profiles.json")